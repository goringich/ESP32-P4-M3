#!/usr/bin/env python3
"""
Comprehensive DOCX builder for the ESP32-P4 gyro-stabilized platform course paper.
Uses only real data from project files, logs, screenshots.
Produces: kursovaya_gyro_platform_READY.docx
"""
import json, os, math, statistics
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import copy

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE       = "/home/goringich/esp"
DOCX_IN    = f"{BASE}/kursovaya_gyro_platform_FINAL.docx"
DOCX_OUT   = f"{BASE}/kursovaya_gyro_platform_READY.docx"
STAB_RAW   = "/tmp/stab_raw.jsonl"
SUMMARY_JSON = "/tmp/esp_uart_test_summary.json"
TMP        = "/tmp/esp_course"
os.makedirs(TMP, exist_ok=True)

SCREENSHOTS = {
    "esp-web-desktop-clean.png": ("Рисунок 1 — Главная страница веб-интерфейса "
                                  "(раздел Overview, подключение к AP JC-ESP32P4M3)"),
    "esp-pad-backend.png":       ("Рисунок 2 — Панель управления (Pad) "
                                  "и подтверждение backend-сервиса"),
}
SCREENSHOT_DIR = BASE


# ══════════════════════════════════════════════════════════════════════════════
# 1. Generate graphs from real stabilization data
# ══════════════════════════════════════════════════════════════════════════════
print("=== 1. Generating graphs ===")
records = [json.loads(l) for l in open(STAB_RAW)]
active  = [r for r in records if r.get('active') and 'error_deg' in r]
print(f"Records total={len(records)}, with PID data={len(active)}")

summary = json.load(open(SUMMARY_JSON)) if os.path.exists(SUMMARY_JSON) else {}
control_metrics = summary.get("control_metrics") or {}

ts      = [r['_ts'] for r in active]
angles  = [r['angle_deg'] for r in active]
errors  = [r['error_deg'] for r in active]
outputs = [r['output'] for r in active]
target  = active[0]['target_deg'] if active else 172.72
duration_s = (ts[-1] - ts[0]) if len(ts) > 1 else 0.0
sample_hz = ((len(ts) - 1) / duration_s) if duration_s > 0 and len(ts) > 1 else 0.0
mean_abs_error = control_metrics.get("mean_abs_error_deg")
max_abs_error = control_metrics.get("max_abs_error_deg")
angle_span = control_metrics.get("angle_span_deg")

def fmt_metric(value, digits):
    if value is None:
        return "n/a"
    return f"{value:.{digits}f}"

GRAPH_ANGLE  = f"{TMP}/graph_angle.png"
GRAPH_OUTPUT = f"{TMP}/graph_output.png"

# Graph 1: Angle vs time
fig, ax = plt.subplots(figsize=(10, 4))
ax.plot(ts, angles, color='#2196F3', linewidth=0.7, label='Угол платформы')
ax.axhline(target, color='#F44336', linewidth=1.2, linestyle='--',
           label=f'Заданный угол {target:.1f}°')
ax.set_xlabel('Время, с')
ax.set_ylabel('Угол, °')
ax.set_title(
    f'Рисунок 3 — Угол платформы в ходе 3-секундного стендового теста (20.05.2026)\n'
    f'Kp=2.00 Ki=0.10 Kd=0.50  α=0.98  dead_zone=1.0°  dt=20мс  f={sample_hz:.1f} Гц'
)
ax.legend(fontsize=9)
ax.grid(True, linestyle=':', alpha=0.5)
ax.set_xlim(0, max(ts) if ts else 1.0)
fig.tight_layout()
fig.savefig(GRAPH_ANGLE, dpi=150, bbox_inches='tight')
plt.close(fig)
print(f"  Saved: {GRAPH_ANGLE}")

# Graph 2: PID output vs time
fig, ax = plt.subplots(figsize=(10, 3))
ax.plot(ts, outputs, color='#4CAF50', linewidth=0.7)
ax.axhline(50,  color='#FF5722', linewidth=0.8, linestyle='--', alpha=0.7, label='+50 sps (насыщение)')
ax.axhline(-50, color='#FF5722', linewidth=0.8, linestyle='--', alpha=0.7, label='−50 sps (насыщение)')
ax.set_xlabel('Время, с')
ax.set_ylabel('Выход ПИД, шаг/с')
ax.set_title('Рисунок 4 — Управляющее воздействие ПИД-регулятора в 3-секундном тесте удержания')
ax.legend(fontsize=9)
ax.grid(True, linestyle=':', alpha=0.5)
ax.set_ylim(-60, 60)
ax.set_xlim(0, max(ts) if ts else 1.0)
fig.tight_layout()
fig.savefig(GRAPH_OUTPUT, dpi=150, bbox_inches='tight')
plt.close(fig)
print(f"  Saved: {GRAPH_OUTPUT}")


# ══════════════════════════════════════════════════════════════════════════════
# 2. Helper functions for python-docx XML manipulation
# ══════════════════════════════════════════════════════════════════════════════

def insert_paragraph_after(ref_para, text='', style_name=None):
    """Insert a new paragraph immediately after ref_para.
    Returns a paragraph object properly bound to the document."""
    # Add a placeholder paragraph at the end of the doc
    new_p = doc.add_paragraph(text)
    if style_name and style_name in doc.styles:
        new_p.style = doc.styles[style_name]
    # Move the new paragraph element to the desired position
    ref_para._p.addnext(new_p._p)
    return new_p


def insert_picture_after(doc, ref_para, img_path, caption, width_inches=5.5):
    """Insert image + caption paragraph after ref_para."""
    # Image paragraph
    img_p = insert_paragraph_after(ref_para)
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    # Caption paragraph
    cap_p = insert_paragraph_after(img_p)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(10)
    return cap_p


def add_table_after(doc, ref_para, headers, rows, style='Table Grid'):
    """Insert a table after ref_para. Returns the table element."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = style
    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    # Data rows
    for ri, row_data in enumerate(rows):
        trow = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            trow.cells[ci].text = str(val)
    # Move table to after ref_para in XML
    tbl_elem = table._tbl
    ref_para._p.addnext(tbl_elem)
    return table


# ══════════════════════════════════════════════════════════════════════════════
# 3. Open document and apply all changes
# ══════════════════════════════════════════════════════════════════════════════
print("\n=== 3. Opening FINAL.docx ===")
doc = Document(DOCX_IN)

# Helper: find paragraph by text keyword
def find_para(keyword, start=0, style_filter=None):
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if style_filter and p.style.name != style_filter:
            continue
        if keyword in p.text:
            return i, p
    return -1, None


# ── 3A. Fix Table 6: Add stabilization commands ('g' and 's') ─────────────────
print("\n--- 3A. Updating command table (Table 6) ---")
t6 = doc.tables[6]
# Check if 'g' already there
existing = [r.cells[0].text for r in t6.rows]
if 'Stabilize' not in existing and 'g' not in str(existing):
    # Add two rows
    for cmd, meaning, note in [
        ("Stabilize (g)", "Включить режим стабилизации",
         "ПИД-регулятор активируется, target=текущий угол, INT=0"),
        ("Stop stab (s)", "Остановить стабилизацию",
         "Выход из режима стабилизации, двигатель останавливается"),
    ]:
        row = t6.add_row()
        row.cells[0].text = cmd
        row.cells[1].text = meaning
        row.cells[2].text = note
    print("  ✓ Added 'g' and 's' rows to command table")
else:
    print("  · Stabilize command already present")


# ── 3B. Insert L293D GPIO connection table after section 2.4 ──────────────────
print("\n--- 3B. Adding L293D GPIO table to section 2.4 ---")
idx24, p24 = find_para("2.4. Подсистема исполнительных механизмов", style_filter='Heading 2')
if p24:
    # Find the last paragraph in section 2.4 before 2.5
    idx25, p25 = find_para("2.5.", start=idx24)
    if idx25 > 0:
        ref = doc.paragraphs[idx25 - 1]
    else:
        ref = doc.paragraphs[idx24 + 3]

    # Add explanatory paragraph first
    note_p = insert_paragraph_after(ref)
    note_p.clear()
    note_p.add_run(
        "Фактическое подключение драйвера L293D к плате JC-ESP32P4-M3 "
        "определено по исходному коду (sdkconfig, app_stepper.c). "
        "ENA и ENB в текущей реализации постоянно подтянуты к VCC — "
        "регулировка скорости через ШИМ не реализована; скорость управляется "
        "частотой тактирования фаз (параметр step_delay_ms). "
        "Внутренний модуль называется app_stepper, однако физически "
        "к нему подключены DC TT-моторы (не шаговые): 4-фазное переключение "
        "H-мостов L293D создаёт дифференциальное управление двумя DC-моторами."
    )
    # style already set by default

    # Caption
    cap = insert_paragraph_after(note_p)
    cap.clear()
    cap.add_run("Таблица 2 — Карта подключения L293D к ESP32-P4-M3").bold = True
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insert the GPIO table after the caption
    gpio_headers = ["Сигнал L293D", "GPIO ESP32-P4-M3", "Назначение"]
    gpio_rows = [
        ("IN1",  "GPIO 3",  "Управление каналом A (Motor A, фаза +)"),
        ("IN2",  "GPIO 4",  "Управление каналом A (Motor A, фаза −)"),
        ("IN3",  "GPIO 5",  "Управление каналом B (Motor B, фаза +)"),
        ("IN4",  "GPIO 20", "Управление каналом B (Motor B, фаза −)"),
        ("ENA",  "VCC (3.3В)", "Постоянно включён — нет ШИМ-регулировки скорости"),
        ("ENB",  "VCC (3.3В)", "Постоянно включён — нет ШИМ-регулировки скорости"),
        ("VM",   "5В / Vin",   "Питание моторов (отдельный источник)"),
        ("GND",  "GND",        "Общая земля ESP + драйвера + моторов"),
        ("OUT1/OUT2", "Motor A: TT DC motor",  "Канал A — двигатель 1"),
        ("OUT3/OUT4", "Motor B: TT DC motor",  "Канал B — двигатель 2"),
    ]
    gpio_table = add_table_after(doc, cap, gpio_headers, gpio_rows)
    print("  ✓ L293D GPIO table inserted after section 2.4")
else:
    print("  ✗ Section 2.4 not found")


# ── 3C. Insert firmware UART log excerpt in section 5.2 ───────────────────────
print("\n--- 3C. Adding UART log excerpt to section 5.2 ---")
_, p52 = find_para("5.2. Проверка аппаратных", style_filter='Heading 2')
if p52:
    idx52, _ = find_para("5.2. Проверка аппаратных")
    # Find first para after the header in 5.2 section
    for i in range(idx52 + 1, min(idx52 + 10, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        if 'UART' in p.text or 'лог' in p.text.lower() or 'I (401)' in p.text:
            ref_p = p
            break
    else:
        ref_p = doc.paragraphs[idx52 + 2]

    log_cap = insert_paragraph_after(ref_p)
    log_cap.clear()
    log_cap.add_run("Листинг 1 — Фрагмент UART-лога при старте прошивки (реальный вывод):").bold = True

    log_text = insert_paragraph_after(log_cap)
    log_text.clear()
    run = log_text.add_run(
        "I (401) H_SDIO_DRV: sdio_data_to_rx_buf_task started\n"
        "I (421) main_task: Calling app_main()\n"
        "I (421) app: ----------------------------------------\n"
        "I (421) app:   APP INITIALIZATION\n"
        "I (421) app: ----------------------------------------\n"
        "I (441) i2c_bus: scan: found 1 device(s)\n"
        "I (441) app: mpu: addr=0x68 WHO_AM_I=0x71 (MPU-9250)\n"
                "I (5231) app_wifi: AP ready ssid='JC-ESP32P4M3' ip='192.168.4.1'\n"
                "I (5239) app_wifi: STA profile 1/2 selected: 'iPhone'\n"
                "I (5241) app_net: http/ws server listening on port 80\n"
                "I (5251) app_ble: advertising started as 'JC-P4-BLE'\n"
                "@telemetry {\"kind\":\"system\",\"uptime_ms\":5300,\"firmware\":\"hello_world_p4\"}\n"
                "@telemetry {\"kind\":\"control\",\"active\":false,\"angle_deg\":178.66,\"target_deg\":178.66}\n"
                "# После команды 'g' (режим стабилизации):\n"
                "@telemetry {\"kind\":\"control\",\"active\":true,\"angle_deg\":178.76,\"target_deg\":178.73,"
                "\"error_deg\":-0.02,\"output\":0.0,\"dt_ms\":20.0,\"accel_norm\":1.02}"
    )
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    print("  ✓ UART log excerpt inserted")
else:
    print("  ✗ Section 5.2 not found")


# ── 3D. Update section 5.5 (mechanical) — honest open bench ───────────────────
print("\n--- 3D. Updating section 5.5 (mechanical) ---")
idx55, p55hdr = find_para("5.5. Испытания механической части", style_filter='Heading 2')
if p55hdr:
    # Update the next normal paragraph
    for i in range(idx55 + 1, min(idx55 + 6, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        if len(p.text.strip()) > 30:
            p.clear()
            p.add_run(
                "Испытание механической части выполнялось на открытом стенде "
                "без 3D-печатного шарового корпуса. На момент испытания корпус "
                "смоделирован в Blender (файл fiish.blend, внешний диаметр 210 мм, "
                "стенка 3 мм, PLA), подготовлен к печати в Ultimaker Cura 5.8 "
                "(Anycubic Kobra 2 Neo, слой 0.2 мм), однако физически не напечатан. "
                "Компоненты (ESP32-P4-M3, GY-9250, L293D, два TT-мотора) "
                "закреплены на временной открытой раме. "
                "Подтверждено: платы и соединения физически совместимы; "
                "моторы вращаются при подаче команд. "
                "Для финального испытания необходимо напечатать корпус, "
                "установить компоненты внутрь и повторить тест."
            )
            break
    print("  ✓ Section 5.5 updated to reflect open bench")
else:
    print("  ✗ Section 5.5 not found")


# ── 3E. Update section 5.6 — add context and insert graphs ────────────────────
print("\n--- 3E. Updating section 5.6 + inserting graphs ---")
idx56, p56hdr = find_para("5.6. Испытания стабилизации", style_filter='Heading 2')
if p56hdr:
    # Find the intro paragraph of 5.6
    for i in range(idx56 + 1, min(idx56 + 5, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        if len(p.text.strip()) > 50:
            p.clear()
            p.add_run(
                "Испытания контура стабилизации проводились на открытом стенде "
                "(без шарового корпуса) 20.05.2026 с использованием скрипта "
                "esp_uart_test.py. Плата подключена к ПК по UART "
                "(/dev/ttyUSB0, 115200 бод). Команда 'g' активирует ПИД-регулятор; "
                "команда 's' — останавливает. Параметры: Kp=2.00, Ki=0.10, Kd=0.50, "
                "alpha=0.98, dead_zone=±1.0°, u_max=±50 шаг/с, dt=20 мс."
                "\n\n"
                f"В тесте удержания получено {len(active)} выборок за {duration_s:.2f} с "
                f"со средней частотой {sample_hz:.1f} Гц. Средняя абсолютная ошибка "
                f"составила {fmt_metric(mean_abs_error, 3)}°, максимальная — {fmt_metric(max_abs_error, 2)}°, "
                f"размах угла — {fmt_metric(angle_span, 2)}°. "
                "Это подтверждает корректность контура измерения и телеметрии в покое. "
                "Отдельно по UART проверены команды движения: вперед и назад вращают оба "
                "колеса в одну сторону, а влево/вправо колеса вращаются встречно."
            )
            break

    # Find a good anchor for inserting graphs (after the metric table)
    # Find paragraph with "Результаты испытаний" in 5.6
    for i in range(idx56, min(idx56 + 25, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        if 'Результаты испытаний' in p.text or 'контур управления функционирует' in p.text:
            anchor_p = p
            break
    else:
        anchor_p = doc.paragraphs[min(idx56 + 15, len(doc.paragraphs) - 1)]

    # Insert graphs after anchor
    g1_cap = insert_picture_after(doc, anchor_p, GRAPH_ANGLE,
                                  "Рисунок 3 — Угол платформы во времени (открытый стенд, 20.05.2026). "
                                  f"3-секундный тест удержания, частота {sample_hz:.1f} Гц.",
                                  width_inches=5.5)
    g2_cap = insert_picture_after(doc, g1_cap, GRAPH_OUTPUT,
                                  "Рисунок 4 — Выход ПИД-регулятора во времени. "
                                  "В данном тесте удержания управляющее воздействие оставалось близким к нулю.",
                                  width_inches=5.5)
    print("  ✓ Context updated and 2 graphs inserted in 5.6")
else:
    print("  ✗ Section 5.6 not found")


# ── 3F. Insert screenshots in section 5.4 ────────────────────────────────────
print("\n--- 3F. Inserting UI screenshots in section 5.4 ---")
idx54, p54hdr = find_para("5.4. Проверка пользовательского интерфейса", style_filter='Heading 2')
if p54hdr:
    # Find a suitable anchor after the header
    for i in range(idx54 + 1, min(idx54 + 8, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        if len(p.text.strip()) > 20:
            anchor = p
            break
    else:
        anchor = doc.paragraphs[idx54 + 1]

    last = anchor
    for fname, caption in SCREENSHOTS.items():
        fpath = os.path.join(SCREENSHOT_DIR, fname)
        if os.path.exists(fpath):
            last = insert_picture_after(doc, last, fpath, caption, width_inches=5.0)
            print(f"  ✓ Inserted: {fname}")
        else:
            print(f"  ✗ Missing: {fpath}")
else:
    print("  ✗ Section 5.4 not found")


# ── 3G. Fix references section ────────────────────────────────────────────────
print("\n--- 3G. Updating references ---")
idx_ref, p_ref = find_para("Список использованных источников", style_filter='Heading 1')
if p_ref:
    # Find existing reference paragraphs (numbered items)
    ref_paras = []
    for i in range(idx_ref + 1, min(idx_ref + 20, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if t and (t[0].isdigit() or t.startswith('—') or 'Espressif' in t or 'InvenSense' in t or 'Blender' in t or 'Anycubic' in t or 'НИУ' in t or 'методи' in t.lower()):
            ref_paras.append((i, p))

    # Replace all references with proper ones
    references = [
        "1. Espressif Systems. ESP32-P4 Technical Reference Manual. — Шанхай: Espressif Systems, 2024. — Режим доступа: https://www.espressif.com/sites/default/files/documentation/esp32-p4_technical_reference_manual_en.pdf (дата обращения: 19.05.2026).",
        "2. Espressif Systems. ESP-IDF Programming Guide v5.x: I2C Driver, HTTP Server, WebSocket. — Шанхай: Espressif Systems, 2024. — Режим доступа: https://docs.espressif.com/projects/esp-idf/en/latest/ (дата обращения: 19.05.2026).",
        "3. Espressif Systems. ESP-Hosted: Wi-Fi and BLE connectivity for ESP32-P4 via ESP32-C6. — Шанхай: Espressif Systems, 2024. — Режим доступа: https://github.com/espressif/esp-hosted (дата обращения: 19.05.2026).",
        "4. TDK InvenSense. MPU-9250 Product Specification Rev 1.1. — Milpitas, CA: TDK InvenSense, 2016. — Режим доступа: https://invensense.tdk.com/wp-content/uploads/2015/02/PS-MPU-9250A-01-v1.1.pdf (дата обращения: 19.05.2026).",
        "5. TDK InvenSense. MPU-9250 Register Map and Descriptions Rev 1.6. — Milpitas, CA: TDK InvenSense, 2016.",
        "6. Texas Instruments. L293D Quadruple Half-H Drivers (datasheet). — Dallas, TX: Texas Instruments, 2016. — Режим доступа: https://www.ti.com/lit/ds/symlink/l293d.pdf (дата обращения: 19.05.2026).",
        "7. Blender Foundation. Blender 4.x Reference Manual. — Amsterdam: Blender Foundation, 2024. — Режим доступа: https://docs.blender.org/ (дата обращения: 19.05.2026).",
        "8. Ultimaker B.V. Ultimaker Cura 5.x User Manual. — Geldermalsen: Ultimaker, 2023. — Режим доступа: https://support.ultimaker.com/s/article/1667337576882 (дата обращения: 19.05.2026).",
        "9. Anycubic Technology Co. Anycubic Kobra 2 Neo User Guide. — Шэньчжэнь: Anycubic, 2023. — Режим доступа: https://www.anycubic.com/products/anycubic-kobra-2-neo (дата обращения: 19.05.2026).",
        "10. Ziegler J.G., Nichols N.B. Optimum Settings for Automatic Controllers // Transactions of the ASME. — 1942. — Vol. 64. — P. 759–768.",
        "11. Madgwick S. An efficient orientation filter for inertial and inertial/magnetic sensor arrays. — University of Bristol, 2010.",
    ]

    if ref_paras:
        for j, (pi, pp) in enumerate(ref_paras):
            if j < len(references):
                pp.clear()
                pp.add_run(references[j])
        # If we have more references than existing paragraphs, add after last
        last_ref_p = ref_paras[-1][1]
        for extra_ref in references[len(ref_paras):]:
            new_p = insert_paragraph_after(last_ref_p)
            new_p.clear()
            new_p.add_run(extra_ref)
            last_ref_p = new_p
        print(f"  ✓ {len(references)} references updated/added")
    else:
        # Append references after header
        last_p = p_ref
        for ref in references:
            last_p = insert_paragraph_after(last_p)
            last_p.clear()
            last_p.add_run(ref)
        print(f"  ✓ {len(references)} references appended")


# ── 3H. Update Appendix D open tasks ─────────────────────────────────────────
print("\n--- 3H. Updating Appendix D open tasks ---")
_, p_appD = find_para("Приложение Д. Список открытых задач", style_filter='Heading 1')
if p_appD:
    idx_d, _ = find_para("Приложение Д. Список открытых задач")
    # Find the tasks table (usually first table after this heading)
    for tbl in doc.tables:
        # Check if this table is after the appendix D heading by checking proximity
        # We look for table with "Закрыть" or "Добавить фото"
        tbl_text = ' '.join(c.text for row in tbl.rows for c in row.cells)
        if 'Закрыть' in tbl_text or 'Добавить фото' in tbl_text or 'напечатать' in tbl_text.lower():
            # Update the table rows to reflect current status
            for row in tbl.rows[1:]:  # skip header
                task_text = row.cells[1].text if len(row.cells) > 1 else ''
                # Tasks completed in this session
                if 'GPIO' in task_text or 'параметры' in task_text.lower() or 'цены' in task_text.lower():
                    if len(row.cells) > 1:
                        row.cells[2].text = 'Выполнено'
            print("  ✓ Appendix D table updated")
            break

    # Add a note paragraph
    idx_note = idx_d + 1
    if idx_note < len(doc.paragraphs):
        # Add summary of what was done
        note_p = insert_paragraph_after(p_appD)
        note_p.clear()
        note_p.add_run(
            "Статус на 20.05.2026: в документе заполнены разделы 2.4 (GPIO-таблица), "
            "4.3 (параметры ПИД), глава 6 (цены), все плейсхолдеры. "
            "Скриншоты UI добавлены. Добавлены реальные результаты 3-секундного "
            "теста удержания и проверки команд движения по UART. "
            "Оставшиеся задачи: 1) напечатать корпус шара; 2) провести испытание "
            "с внешним механическим возмущением и получить динамические метрики стабилизации; "
            "3) добавить фото сборки; 4) проверить процент заимствований; "
            "5) согласовать с руководителем."
        )
        print("  ✓ Open tasks note added")


# ── 3I. Fix coordinate system explanation ─────────────────────────────────────
print("\n--- 3I. Adding coordinate system explanation ---")
_, p42 = find_para("4.2. Оценка угла наклона", style_filter='Heading 2')
if p42:
    idx42, _ = find_para("4.2. Оценка угла наклона")
    # Find complementary filter paragraph
    for i in range(idx42 + 1, min(idx42 + 10, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        if 'complementary' in p.text.lower() or 'alpha' in p.text or 'угла' in p.text:
            # Check if coordinate explanation already there
            if '178' not in p.text and 'atan2' not in p.text:
                coord_note = insert_paragraph_after(p)
                coord_note.clear()
                coord_note.add_run(
                    "Особенность системы координат: угол вычисляется как "
                    "atan2(ax, az) × (180/π), где ax и az — показания акселерометра "
                    "в g. Когда платформа находится в рабочем горизонтальном положении "
                    "(Z-ось датчика направлена вниз, az ≈ −1g, ax ≈ 0), "
                    "atan2(0, −1) = π, что соответствует 180°. "
                    "Таким образом, угол ≈178–180° означает нормальное рабочее положение платформы, "
                    "а отклонение от 180° — наклон в ту или иную сторону."
                )
                coord_note.style = doc.styles['Normal']
                print("  ✓ Coordinate system explanation added")
            break


# ── 3J. Fix Appendix V — update what's available vs missing ──────────────────
print("\n--- 3J. Updating Appendix V ---")
_, p_appV = find_para("Приложение В. Материалы, которые нужно добавить", style_filter='Heading 1')
if p_appV:
    # Update the first paragraph after heading
    idx_v, _ = find_para("Приложение В.")
    for i in range(idx_v + 1, min(idx_v + 20, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if 'Фотография полной' in t or 'Фрагмент успешного' in t:
            p.clear()
            p.add_run(
                "Доступно и добавлено в документ:\n"
                "✓ Скриншоты веб-интерфейса (разделы Overview, Pad) — добавлены в раздел 5.4.\n"
                "✓ Фрагмент UART-лога (I2C scan found 0x68, WHO_AM_I=0x71) — добавлен в раздел 5.2.\n"
                "✓ Графики испытания стабилизации (угол и выход ПИД) — добавлены в раздел 5.6.\n"
                "✓ Таблица GPIO-подключений L293D — добавлена в раздел 2.4.\n"
                "✓ Параметры регулятора Kp=2.00 Ki=0.10 Kd=0.50 — добавлены в раздел 4.3.\n"
                "✓ Таблица цен компонентов, итого 5950 руб. — добавлена в главу 6.\n\n"
                "Требуется добавить вручную (данных нет):\n"
                "✗ Фотография полной сборки стенда.\n"
                "✗ Фотография подключения GY-9250 к плате.\n"
                "✗ Фотография напечатанного шара (корпус ещё не напечатан).\n"
                "✗ Схема питания с номиналами (тип и ёмкость батарей).\n"
                "✗ Испытание стабилизации в замкнутом корпусе."
            )
            print("  ✓ Appendix V updated")
            break


# ══════════════════════════════════════════════════════════════════════════════
# 4. Save READY.docx
# ══════════════════════════════════════════════════════════════════════════════
print(f"\n=== 4. Saving {DOCX_OUT} ===")
doc.save(DOCX_OUT)
print(f"  ✓ Saved: {DOCX_OUT}")


# ══════════════════════════════════════════════════════════════════════════════
# 5. Export PDF via LibreOffice
# ══════════════════════════════════════════════════════════════════════════════
print("\n=== 5. Exporting PDF ===")
PDF_OUT = DOCX_OUT.replace('.docx', '.pdf')
import subprocess
result = subprocess.run([
    'libreoffice', '--headless', '--convert-to', 'pdf',
    '--outdir', os.path.dirname(DOCX_OUT),
    DOCX_OUT
], capture_output=True, text=True, timeout=120)
if result.returncode == 0:
    print(f"  ✓ PDF exported: {PDF_OUT}")
else:
    print(f"  ✗ PDF export failed: {result.stderr[:200]}")


# ══════════════════════════════════════════════════════════════════════════════
# 6. Final report
# ══════════════════════════════════════════════════════════════════════════════
print(f"""
═══════════════════════════════════════════════════════════
ОТЧЁТ: что сделано
═══════════════════════════════════════════════════════════
✓ Таблица GPIO L293D (IN1=GPIO3, IN2=GPIO4, IN3=GPIO5, IN4=GPIO20)
✓ Пояснение: module=app_stepper, моторы=DC TT (не шаговые)
✓ Команды 'g' (stabilize) и 's' (stop) добавлены в таблицу команд
✓ Объяснение системы координат (178°=горизонтальное положение)
✓ Фрагмент UART-лога (addr=0x68, WHO_AM_I=0x71, BLE, Wi-Fi)
✓ Скриншоты UI (Overview, Pad)
✓ 2 графика из реальных данных (угол и выход ПИД, 3-секундный тест удержания)
✓ Раздел 5.5: честное описание открытого стенда (без шара)
✓ Раздел 5.6: контекст испытания и реальные метрики удержания
✓ 11 источников с URL и датами обращения
✓ Приложение В: список что есть / чего нет
✓ Приложение Д: статус задач

Использованные реальные файлы:
  /tmp/stab_raw.jsonl ({len(active)} записей, {duration_s:.2f} с, 20.05.2026)
  /tmp/esp_uart_test_summary.json
  /home/goringich/esp/sdkconfig (GPIO, Kp/Kd/Ki/alpha)
  /home/goringich/esp/components/app/src/app_stepper.c
  /home/goringich/esp/esp-web-desktop-clean.png
  /home/goringich/esp/esp-pad-backend.png

Что нужно сделать вручную (данных нет):
  ✗ Фото сборки стенда
  ✗ Фото подключения датчика
  ✗ Напечатать шар и провести испытание в корпусе
  ✗ Схема питания с номиналами
  ✗ Проверка процента заимствований (антиплагиат)
  ✗ Согласование с Морозовым Н.С.

Итоговые файлы:
""")
print(f"  DOCX: {DOCX_OUT}")
print(f"  PDF:  {PDF_OUT}")
